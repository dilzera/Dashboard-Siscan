"""
Script para gerar documento Word de especificação do sistema
Central Inteligente de Câncer de Mama - Curitiba
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_specification_document():
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Heading 1']
    style.font.color.rgb = RGBColor(0x17, 0xa2, 0xb8)
    style.font.size = Pt(18)
    
    style2 = doc.styles['Heading 2']
    style2.font.color.rgb = RGBColor(0x13, 0x84, 0x96)
    style2.font.size = Pt(14)
    
    style3 = doc.styles['Heading 3']
    style3.font.color.rgb = RGBColor(0x34, 0x3a, 0x40)
    style3.font.size = Pt(12)
    
    # CAPA
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title_run = title.add_run("CENTRAL INTELIGENTE DE CÂNCER DE MAMA")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x17, 0xa2, 0xb8)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("CURITIBA")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(0xff, 0x69, 0xb4)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc_type = doc.add_paragraph()
    doc_type_run = doc_type.add_run("DOCUMENTO DE ESPECIFICAÇÃO FUNCIONAL")
    doc_type_run.bold = True
    doc_type_run.font.size = Pt(16)
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.add_run(f"Versão: 1.0\n")
    info.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}\n")
    info.add_run("Sistema: Dashboard SISCAN\n")
    info.add_run("Secretaria Municipal de Saúde de Curitiba")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # SUMÁRIO
    doc.add_heading("SUMÁRIO", level=1)
    sumario = [
        "1. Visão Geral do Sistema",
        "2. Arquitetura Técnica",
        "3. Autenticação e Segurança",
        "4. Tela Principal - Painel de Desempenho",
        "5. Aba de Auditoria de Risco",
        "6. Aba de Auditoria de Outliers",
        "7. Aba de Auditoria de Qualidade de Dados",
        "8. Aba de Indicadores",
        "9. Aba de Análise por Unidade de Saúde",
        "10. Sistema de Priorização de Pacientes",
        "11. Aba de Dados de Pacientes",
        "12. Aba de Navegação de Pacientes",
        "13. Aba de Interoperabilidade (eSaúde)",
        "14. Regras de Negócio",
        "15. Filtros Globais",
        "16. Glossário"
    ]
    for item in sumario:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # 1. VISÃO GERAL
    doc.add_heading("1. Visão Geral do Sistema", level=1)
    
    doc.add_heading("1.1 Objetivo", level=2)
    doc.add_paragraph(
        "A Central Inteligente de Câncer de Mama de Curitiba é um sistema de dashboard interativo "
        "desenvolvido para monitoramento e análise de dados de mamografia do sistema SISCAN "
        "(Sistema de Informação do Câncer). O sistema visa auxiliar gestores e profissionais de saúde "
        "na tomada de decisões, identificação de casos prioritários e melhoria contínua dos processos "
        "de rastreamento e diagnóstico do câncer de mama."
    )
    
    doc.add_heading("1.2 Funcionalidades Principais", level=2)
    funcionalidades = [
        "Visualização de KPIs (Indicadores-Chave de Desempenho) em tempo real",
        "Análise de desempenho por unidade de saúde e distrito sanitário",
        "Auditoria de riscos baseada em classificação BI-RADS",
        "Detecção de outliers e inconsistências nos dados",
        "Sistema de priorização de pacientes com algoritmo baseado em risco",
        "Navegação de pacientes com múltiplos atendimentos",
        "Interoperabilidade com sistema eSaúde",
        "Exportação de dados para busca ativa"
    ]
    for func in funcionalidades:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_heading("1.3 Público-Alvo", level=2)
    doc.add_paragraph(
        "O sistema é destinado a gestores de saúde, coordenadores de programas de rastreamento, "
        "profissionais de saúde das unidades, auditores e equipes de qualidade da Secretaria "
        "Municipal de Saúde de Curitiba."
    )
    
    doc.add_page_break()
    
    # 2. ARQUITETURA TÉCNICA
    doc.add_heading("2. Arquitetura Técnica", level=1)
    
    doc.add_heading("2.1 Tecnologias Utilizadas", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Tecnologia'
    
    techs = [
        ("Backend", "Python 3.11, Flask"),
        ("Frontend", "Dash 2.18.2, Dash Bootstrap Components 1.7.1"),
        ("Visualizações", "Plotly"),
        ("Banco de Dados", "PostgreSQL"),
        ("ORM", "SQLAlchemy"),
        ("Servidor Web", "Gunicorn"),
        ("Autenticação", "Flask-Login")
    ]
    for comp, tech in techs:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech
    
    doc.add_paragraph()
    
    doc.add_heading("2.2 Estrutura de Dados Principal", level=2)
    doc.add_paragraph("Tabela: exam_records")
    
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Campo'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Descrição'
    
    campos = [
        ("id", "Integer", "Identificador único do registro"),
        ("cartao_sus", "String", "Número do Cartão SUS do paciente"),
        ("paciente__nome", "String", "Nome do paciente"),
        ("unidade_de_saude__nome", "String", "Nome da unidade de saúde"),
        ("unidade_de_saude__distrito", "String", "Distrito sanitário"),
        ("unidade_de_saude__data_da_solicitacao", "Date", "Data da solicitação do exame"),
        ("prestador_de_servico__data_da_realizacao", "Date", "Data de realização do exame"),
        ("data_liberacao", "Date", "Data de liberação do resultado"),
        ("birads_max", "String", "Classificação BI-RADS máxima"),
        ("wait_days", "Integer", "Dias de espera entre solicitação e realização"),
        ("idade", "Integer", "Idade do paciente")
    ]
    for campo, tipo, desc in campos:
        row = table2.add_row().cells
        row[0].text = campo
        row[1].text = tipo
        row[2].text = desc
    
    doc.add_page_break()
    
    # 3. AUTENTICAÇÃO E SEGURANÇA
    doc.add_heading("3. Autenticação e Segurança", level=1)
    
    doc.add_heading("3.1 Tela de Login", level=2)
    doc.add_paragraph(
        "O acesso ao sistema é protegido por autenticação obrigatória. A tela de login apresenta "
        "campos para usuário e senha, com visual alinhado à identidade do sistema (cores ciano e rosa)."
    )
    
    doc.add_heading("3.2 Regras de Autenticação", level=2)
    regras_auth = [
        "Sessão expira automaticamente após 1 hora de inatividade",
        "Senhas são criptografadas usando algoritmo scrypt (werkzeug)",
        "Tentativas de acesso a páginas protegidas redirecionam para login",
        "URL original é preservada após autenticação bem-sucedida",
        "Usuário administrador configurável via variável de ambiente"
    ]
    for regra in regras_auth:
        doc.add_paragraph(regra, style='List Bullet')
    
    doc.add_heading("3.3 Usuários do Sistema", level=2)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Usuário'
    hdr[1].text = 'Perfil'
    hdr[2].text = 'Configuração'
    
    users = [
        ("admin", "Administrador", "Senha via variável ADMIN_PASSWORD"),
        ("neusa", "Usuário padrão", "Senha via variável NEUSA_PASSWORD")
    ]
    for user, perfil, config in users:
        row = table3.add_row().cells
        row[0].text = user
        row[1].text = perfil
        row[2].text = config
    
    doc.add_page_break()
    
    # 4. PAINEL DE DESEMPENHO
    doc.add_heading("4. Tela Principal - Painel de Desempenho", level=1)
    
    doc.add_heading("4.1 Descrição", level=2)
    doc.add_paragraph(
        "A tela principal apresenta uma visão consolidada do desempenho do programa de mamografia, "
        "com KPIs, gráficos de tendência e análises comparativas entre unidades e distritos."
    )
    
    doc.add_heading("4.2 KPIs Exibidos", level=2)
    table_kpi = doc.add_table(rows=1, cols=3)
    table_kpi.style = 'Table Grid'
    hdr = table_kpi.rows[0].cells
    hdr[0].text = 'KPI'
    hdr[1].text = 'Descrição'
    hdr[2].text = 'Fórmula'
    
    kpis = [
        ("Total de Exames", "Quantidade total de mamografias no período", "COUNT(registros)"),
        ("Tempo Médio de Espera", "Média de dias entre solicitação e realização", "AVG(wait_days)"),
        ("Tempo Mediano de Espera", "Mediana de dias de espera", "MEDIAN(wait_days)"),
        ("Taxa de Conformidade", "Percentual de exames dentro do prazo (30 dias)", "(wait_days <= 30) / total * 100"),
        ("Casos Alto Risco", "Quantidade de exames BI-RADS 4 ou 5", "COUNT(birads IN (4,5))")
    ]
    for kpi, desc, formula in kpis:
        row = table_kpi.add_row().cells
        row[0].text = kpi
        row[1].text = desc
        row[2].text = formula
    
    doc.add_paragraph()
    
    doc.add_heading("4.3 Visualizações", level=2)
    
    doc.add_heading("4.3.1 Gráfico de Volume Mensal", level=3)
    doc.add_paragraph(
        "Gráfico de linhas mostrando a evolução do número de exames realizados por mês. "
        "Permite identificar tendências sazonais e variações no volume de atendimento."
    )
    
    doc.add_heading("4.3.2 Conformidade por Unidade", level=3)
    doc.add_paragraph(
        "Gráfico de barras horizontais ordenado pela taxa de conformidade, permitindo "
        "comparar o desempenho das unidades de saúde no cumprimento da meta de 30 dias."
    )
    
    doc.add_heading("4.3.3 Distribuição por Distrito", level=3)
    doc.add_paragraph(
        "Gráfico de pizza mostrando a proporção de exames por distrito sanitário, "
        "auxiliando na análise de cobertura territorial."
    )
    
    doc.add_page_break()
    
    # 5. AUDITORIA DE RISCO
    doc.add_heading("5. Aba de Auditoria de Risco", level=1)
    
    doc.add_heading("5.1 Objetivo", level=2)
    doc.add_paragraph(
        "Identificar e monitorar casos de alto risco baseados na classificação BI-RADS, "
        "permitindo ações proativas de acompanhamento e busca ativa."
    )
    
    doc.add_heading("5.2 Classificação BI-RADS", level=2)
    table_birads = doc.add_table(rows=1, cols=4)
    table_birads.style = 'Table Grid'
    hdr = table_birads.rows[0].cells
    hdr[0].text = 'Categoria'
    hdr[1].text = 'Descrição'
    hdr[2].text = 'Risco de Malignidade'
    hdr[3].text = 'Conduta'
    
    birads = [
        ("0", "Inconclusivo", "N/A", "Avaliação adicional necessária"),
        ("1", "Negativo", "0%", "Rastreamento de rotina"),
        ("2", "Achado benigno", "0%", "Rastreamento de rotina"),
        ("3", "Provavelmente benigno", "< 2%", "Acompanhamento em 6 meses"),
        ("4", "Suspeita de malignidade", "2-95%", "Biópsia recomendada"),
        ("5", "Altamente sugestivo", "> 95%", "Biópsia obrigatória"),
        ("6", "Malignidade comprovada", "100%", "Tratamento oncológico")
    ]
    for cat, desc, risco, conduta in birads:
        row = table_birads.add_row().cells
        row[0].text = cat
        row[1].text = desc
        row[2].text = risco
        row[3].text = conduta
    
    doc.add_paragraph()
    
    doc.add_heading("5.3 Visualizações", level=2)
    doc.add_paragraph("- Gráfico de pizza com distribuição por categoria BI-RADS")
    doc.add_paragraph("- Cards destacando quantidade de casos BI-RADS 4 e 5")
    doc.add_paragraph("- Tabela de casos alto risco com dados para busca ativa")
    
    doc.add_page_break()
    
    # 6. AUDITORIA DE OUTLIERS
    doc.add_heading("6. Aba de Auditoria de Outliers", level=1)
    
    doc.add_heading("6.1 Objetivo", level=2)
    doc.add_paragraph(
        "Detectar e categorizar inconsistências nos dados que podem indicar problemas "
        "de registro, processos ou qualidade da informação."
    )
    
    doc.add_heading("6.2 Tipos de Outliers Detectados", level=2)
    table_out = doc.add_table(rows=1, cols=3)
    table_out.style = 'Table Grid'
    hdr = table_out.rows[0].cells
    hdr[0].text = 'Tipo'
    hdr[1].text = 'Critério'
    hdr[2].text = 'Possível Causa'
    
    outliers = [
        ("Datas Absurdas", "Data de solicitação anterior a 01/01/2020", "Erro de digitação ou migração"),
        ("Delta Negativo", "Data de realização anterior à solicitação", "Inversão de datas no registro"),
        ("BI-RADS Inválido", "Valor fora do padrão (0-6)", "Erro de codificação"),
        ("Espera Excessiva", "Tempo de espera superior a 365 dias", "Registro tardio ou problema de fluxo")
    ]
    for tipo, criterio, causa in outliers:
        row = table_out.add_row().cells
        row[0].text = tipo
        row[1].text = criterio
        row[2].text = causa
    
    doc.add_paragraph()
    
    doc.add_heading("6.3 Regras de Exclusão", level=2)
    doc.add_paragraph(
        "Os outliers são EXCLUÍDOS dos gráficos de performance principais para não distorcer "
        "as métricas. Porém, ficam disponíveis nesta aba para auditoria e correção."
    )
    
    doc.add_page_break()
    
    # 7. QUALIDADE DE DADOS
    doc.add_heading("7. Aba de Auditoria de Qualidade de Dados", level=1)
    
    doc.add_heading("7.1 Objetivo", level=2)
    doc.add_paragraph(
        "Categorizar e quantificar as inconsistências de dados de forma estruturada, "
        "permitindo ações corretivas direcionadas."
    )
    
    doc.add_heading("7.2 Categorias de Inconsistência", level=2)
    categorias = [
        ("Datas Anteriores a 2020", "Registros com data de solicitação antes de 01/01/2020"),
        ("Delta Negativo", "Data de realização anterior à data de solicitação"),
        ("BI-RADS Inválido", "Classificação fora do padrão estabelecido"),
        ("Espera > 365 dias", "Tempo entre solicitação e realização superior a um ano"),
        ("Dados Faltantes", "Campos obrigatórios não preenchidos")
    ]
    for cat, desc in categorias:
        p = doc.add_paragraph()
        p.add_run(f"{cat}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 8. INDICADORES
    doc.add_heading("8. Aba de Indicadores", level=1)
    
    doc.add_heading("8.1 Descrição", level=2)
    doc.add_paragraph(
        "Apresenta 10 indicadores clínicos distribuídos em 4 blocos temáticos, "
        "alinhados às diretrizes do Ministério da Saúde para rastreamento do câncer de mama."
    )
    
    doc.add_heading("8.2 Bloco 1 - Cobertura da População-Alvo", level=2)
    doc.add_paragraph("Indicadores relacionados à cobertura de mulheres na faixa etária alvo (50-69 anos).")
    
    doc.add_heading("8.3 Bloco 2 - Agilidade no Acesso e Entrega de Resultados", level=2)
    indicadores_b2 = [
        "Taxa de exames realizados em até 30 dias",
        "Tempo médio entre solicitação e realização",
        "Taxa de resultados liberados em tempo hábil"
    ]
    for ind in indicadores_b2:
        doc.add_paragraph(ind, style='List Bullet')
    
    doc.add_heading("8.4 Bloco 3 - Encaminhamentos por Categoria BI-RADS", level=2)
    doc.add_paragraph("Análise da distribuição de resultados por categoria e encaminhamentos adequados.")
    
    doc.add_heading("8.5 Bloco 4 - Casos Especiais / Fora da Faixa Etária", level=2)
    doc.add_paragraph("Monitoramento de exames realizados fora da faixa etária de rastreamento e casos especiais.")
    
    doc.add_page_break()
    
    # 9. ANÁLISE POR UNIDADE
    doc.add_heading("9. Aba de Análise por Unidade de Saúde", level=1)
    
    doc.add_heading("9.1 Descrição", level=2)
    doc.add_paragraph(
        "Permite análise detalhada do desempenho de cada unidade de saúde individualmente, "
        "com KPIs específicos, distribuição demográfica e fila de priorização."
    )
    
    doc.add_heading("9.2 Componentes da Tela", level=2)
    
    doc.add_heading("9.2.1 Seleção de Unidade", level=3)
    doc.add_paragraph(
        "Dropdown para seleção da unidade de saúde, com opção de pesquisa. "
        "Botão 'Analisar Unidade' para carregar os dados."
    )
    
    doc.add_heading("9.2.2 KPIs da Unidade", level=3)
    doc.add_paragraph("Cards mostrando: Total de Exames, Tempo Médio de Espera, Taxa de Conformidade, Casos Alto Risco.")
    
    doc.add_heading("9.2.3 Heatmap Demográfico", level=3)
    doc.add_paragraph(
        "Mapa de calor mostrando a distribuição de pacientes por faixa etária e mês, "
        "permitindo identificar padrões demográficos de atendimento."
    )
    
    doc.add_heading("9.2.4 Distribuição de Tempo de Espera", level=3)
    doc.add_paragraph(
        "Histograma mostrando a distribuição dos tempos de espera, "
        "com destaque para a meta de 30 dias."
    )
    
    doc.add_heading("9.2.5 Tendência Mensal", level=3)
    doc.add_paragraph(
        "Gráfico de linha mostrando a evolução do tempo médio de espera ao longo dos meses, "
        "com linha de referência da meta de 30 dias."
    )
    
    doc.add_heading("9.2.6 Tabela de Pendentes", level=3)
    doc.add_paragraph(
        "Lista de pacientes com BI-RADS 0, 3, 4 ou 5 que necessitam retorno ou acompanhamento, "
        "ordenados por prioridade."
    )
    
    doc.add_page_break()
    
    # 10. SISTEMA DE PRIORIZAÇÃO
    doc.add_heading("10. Sistema de Priorização de Pacientes", level=1)
    
    doc.add_heading("10.1 Objetivo", level=2)
    doc.add_paragraph(
        "Classificar automaticamente os pacientes em níveis de prioridade baseados no risco, "
        "definindo SLAs (Acordos de Nível de Serviço) e ações recomendadas para cada caso."
    )
    
    doc.add_heading("10.2 Níveis de Prioridade", level=2)
    table_prio = doc.add_table(rows=1, cols=5)
    table_prio.style = 'Table Grid'
    hdr = table_prio.rows[0].cells
    hdr[0].text = 'Nível'
    hdr[1].text = 'BI-RADS'
    hdr[2].text = 'SLA'
    hdr[3].text = 'Cor'
    hdr[4].text = 'Ação Recomendada'
    
    prioridades = [
        ("CRÍTICA", "4, 5", "3 dias", "Vermelho", "Busca ativa imediata, encaminhamento oncologia"),
        ("ALTA", "0", "45 dias", "Laranja", "Agendamento exame complementar"),
        ("MÉDIA", "3", "6 meses", "Amarelo", "Acompanhamento semestral"),
        ("MONITORAMENTO", "6", "Contínuo", "Roxo", "Acompanhamento oncológico"),
        ("ROTINA", "1, 2", "24 meses", "Verde", "Rastreamento bienal")
    ]
    for nivel, birads, sla, cor, acao in prioridades:
        row = table_prio.add_row().cells
        row[0].text = nivel
        row[1].text = birads
        row[2].text = sla
        row[3].text = cor
        row[4].text = acao
    
    doc.add_paragraph()
    
    doc.add_heading("10.3 Algoritmo de Priorização", level=2)
    doc.add_paragraph(
        "O algoritmo considera os seguintes fatores para classificação:"
    )
    fatores = [
        "Classificação BI-RADS (fator principal)",
        "Histórico de câncer via APAC (quando disponível via cruzamento com eSaúde)",
        "Idade do paciente (maior peso para idades de maior risco)",
        "Tempo desde a liberação do resultado"
    ]
    for fator in fatores:
        doc.add_paragraph(fator, style='List Bullet')
    
    doc.add_heading("10.4 Regras de Escalação", level=2)
    doc.add_paragraph(
        "Pacientes com histórico de APAC de câncer (identificados via interoperabilidade) "
        "têm sua prioridade automaticamente elevada em um nível."
    )
    
    doc.add_heading("10.5 Cards de Resumo", level=2)
    doc.add_paragraph(
        "A interface exibe cards coloridos com a contagem de pacientes em cada nível de prioridade, "
        "permitindo visão rápida da situação da unidade."
    )
    
    doc.add_heading("10.6 Fila de Priorização", level=2)
    doc.add_paragraph(
        "Tabela ordenada por prioridade mostrando: Nome do paciente, Cartão SUS, BI-RADS, "
        "Data de liberação, Idade, Nível de prioridade, SLA e Ação recomendada."
    )
    
    doc.add_heading("10.7 Exportação para Busca Ativa", level=2)
    doc.add_paragraph(
        "Botão 'Encaminhar para busca ativa' permite exportar lista de pacientes de alto risco "
        "(BI-RADS 4 e 5) em formato CSV para ações de campo."
    )
    
    doc.add_page_break()
    
    # 11. DADOS DE PACIENTES
    doc.add_heading("11. Aba de Dados de Pacientes", level=1)
    
    doc.add_heading("11.1 Descrição", level=2)
    doc.add_paragraph(
        "Listagem completa de todos os registros do sistema com filtros avançados "
        "e paginação para navegação eficiente."
    )
    
    doc.add_heading("11.2 Filtros Disponíveis", level=2)
    filtros = [
        "Nome do paciente (busca parcial)",
        "Cartão SUS",
        "Unidade de Saúde",
        "Distrito Sanitário",
        "Classificação BI-RADS",
        "Período (data inicial e final)"
    ]
    for filtro in filtros:
        doc.add_paragraph(filtro, style='List Bullet')
    
    doc.add_heading("11.3 Colunas Exibidas", level=2)
    doc.add_paragraph(
        "Nome, Cartão SUS, Unidade, Distrito, Data Solicitação, Data Realização, "
        "Dias de Espera, BI-RADS, Idade."
    )
    
    doc.add_page_break()
    
    # 12. NAVEGAÇÃO DE PACIENTES
    doc.add_heading("12. Aba de Navegação de Pacientes", level=1)
    
    doc.add_heading("12.1 Objetivo", level=2)
    doc.add_paragraph(
        "Acompanhar a trajetória de pacientes com múltiplos atendimentos ao longo do tempo, "
        "visualizando o histórico completo de exames e resultados."
    )
    
    doc.add_heading("12.2 Funcionalidades", level=2)
    func_nav = [
        "Busca por nome ou Cartão SUS",
        "Timeline de atendimentos ordenada cronologicamente",
        "Exibição de BI-RADS de cada exame",
        "Unidade de atendimento em cada visita",
        "Tempo de espera em cada atendimento",
        "Identificação de mudanças na classificação BI-RADS"
    ]
    for func in func_nav:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_page_break()
    
    # 13. INTEROPERABILIDADE
    doc.add_heading("13. Aba de Interoperabilidade (eSaúde)", level=1)
    
    doc.add_heading("13.1 Objetivo", level=2)
    doc.add_paragraph(
        "Realizar cruzamento de dados entre o sistema SISCAN e o sistema eSaúde "
        "para enriquecer as informações dos pacientes e identificar históricos relevantes."
    )
    
    doc.add_heading("13.2 Dados Integrados", level=2)
    dados_int = [
        "Histórico de APAC oncológica (indica tratamento prévio de câncer)",
        "Dados cadastrais complementares",
        "Informações de outros atendimentos no eSaúde"
    ]
    for dado in dados_int:
        doc.add_paragraph(dado, style='List Bullet')
    
    doc.add_heading("13.3 Tabela termo_linkage", level=2)
    doc.add_paragraph(
        "Armazena os dados cruzados entre sistemas, vinculados pelo Cartão SUS do paciente."
    )
    
    doc.add_heading("13.4 Cards de Resumo", level=2)
    doc.add_paragraph(
        "Exibe estatísticas do cruzamento: total de registros vinculados, "
        "pacientes com histórico de APAC, pacientes sem vínculo."
    )
    
    doc.add_heading("13.5 Busca e Filtros", level=2)
    doc.add_paragraph(
        "Permite buscar pacientes específicos e filtrar por status de vínculo "
        "e presença de histórico oncológico."
    )
    
    doc.add_page_break()
    
    # 14. REGRAS DE NEGÓCIO
    doc.add_heading("14. Regras de Negócio", level=1)
    
    doc.add_heading("14.1 Período de Análise", level=2)
    doc.add_paragraph(
        "O sistema considera apenas dados a partir de 01/01/2023 para métricas de performance. "
        "Dados anteriores são considerados outliers e disponibilizados apenas na aba de auditoria."
    )
    
    doc.add_heading("14.2 Meta de Tempo de Espera", level=2)
    doc.add_paragraph(
        "A meta estabelecida é de 30 dias entre a solicitação e a realização do exame. "
        "Exames realizados dentro deste prazo são considerados 'conformes'."
    )
    
    doc.add_heading("14.3 Exclusão de Outliers", level=2)
    doc.add_paragraph(
        "Para cálculo de métricas de performance, são excluídos automaticamente:"
    )
    exclusoes = [
        "Registros com data de solicitação anterior a 01/01/2020",
        "Registros com tempo de espera negativo",
        "Registros com tempo de espera superior a 365 dias",
        "Registros com BI-RADS inválido ou não preenchido"
    ]
    for exc in exclusoes:
        doc.add_paragraph(exc, style='List Bullet')
    
    doc.add_heading("14.4 Faixa Etária de Rastreamento", level=2)
    doc.add_paragraph(
        "A faixa etária prioritária para rastreamento é de 50 a 69 anos, conforme "
        "diretrizes do Ministério da Saúde. Exames fora desta faixa são monitorados "
        "como casos especiais."
    )
    
    doc.add_heading("14.5 Atualização de Dados", level=2)
    doc.add_paragraph(
        "O botão 'Atualizar Dados' permite recarregar os dados do banco de dados "
        "sob demanda, refletindo as últimas inserções ou correções."
    )
    
    doc.add_page_break()
    
    # 15. FILTROS GLOBAIS
    doc.add_heading("15. Filtros Globais", level=1)
    
    doc.add_heading("15.1 Descrição", level=2)
    doc.add_paragraph(
        "O sistema oferece filtros globais no topo da interface que afetam "
        "todas as visualizações e métricas exibidas."
    )
    
    doc.add_heading("15.2 Filtros Disponíveis", level=2)
    table_filtros = doc.add_table(rows=1, cols=3)
    table_filtros.style = 'Table Grid'
    hdr = table_filtros.rows[0].cells
    hdr[0].text = 'Filtro'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Descrição'
    
    filtros_glob = [
        ("Ano", "Dropdown", "Filtra por ano de solicitação (2023, 2024, 2025)"),
        ("Unidade de Saúde", "Dropdown pesquisável", "Filtra por unidade específica"),
        ("Distrito Sanitário", "Dropdown", "Filtra por distrito (região)"),
        ("Status de Conformidade", "Dropdown", "Filtra por conforme/não conforme")
    ]
    for filtro, tipo, desc in filtros_glob:
        row = table_filtros.add_row().cells
        row[0].text = filtro
        row[1].text = tipo
        row[2].text = desc
    
    doc.add_page_break()
    
    # 16. GLOSSÁRIO
    doc.add_heading("16. Glossário", level=1)
    
    termos = [
        ("BI-RADS", "Breast Imaging Reporting and Data System - Sistema padronizado de classificação de achados mamográficos"),
        ("SISCAN", "Sistema de Informação do Câncer - sistema nacional de registro de exames oncológicos"),
        ("eSaúde", "Sistema de informação em saúde utilizado pela rede municipal de Curitiba"),
        ("APAC", "Autorização de Procedimento de Alta Complexidade - registro de tratamentos oncológicos"),
        ("SLA", "Service Level Agreement - Acordo de Nível de Serviço, prazo estabelecido para ação"),
        ("KPI", "Key Performance Indicator - Indicador-Chave de Desempenho"),
        ("Cartão SUS", "Cartão Nacional de Saúde - identificador único do cidadão no SUS"),
        ("Distrito Sanitário", "Divisão administrativa territorial da saúde municipal"),
        ("Conformidade", "Atendimento dentro do prazo estabelecido (30 dias)"),
        ("Outlier", "Valor atípico ou inconsistente nos dados"),
        ("Busca Ativa", "Ação de contato proativo com pacientes para agendamento ou acompanhamento")
    ]
    
    for termo, definicao in termos:
        p = doc.add_paragraph()
        p.add_run(f"{termo}: ").bold = True
        p.add_run(definicao)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Rodapé
    footer = doc.add_paragraph()
    footer.add_run("Central Inteligente de Câncer de Mama - Curitiba").italic = True
    footer.add_run(f"\nDocumento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    filename = "Especificacao_Central_Inteligente_Cancer_Mama.docx"
    doc.save(filename)
    print(f"Documento salvo: {filename}")
    return filename

if __name__ == "__main__":
    create_specification_document()
